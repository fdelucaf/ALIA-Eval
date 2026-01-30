#!/usr/bin/env python3
"""
Extract parallel paragraphs from DOCX files.
Uses double newline as paragraph separator.
"""

from pathlib import Path
from docx import Document
import re
from typing import List, Dict, Tuple


class DocxParagraphExtractor:
    """Extract paragraphs from DOCX files."""
    
    def __init__(self, docx_path: Path):
        self.docx_path = docx_path
        self.doc = Document(docx_path)
        
    def extract_paragraphs(self) -> Tuple[List[str], Dict[str, int]]:
        """
        Extract paragraphs from the document.
        Returns: (list of paragraphs, dict of issues found)
        """
        paragraphs = []
        issues = {
            'images': 0,
            'tables': 0,
            'headers': 0,
            'footers': 0
        }
        
        # Count non-text elements
        issues['images'] = len([rel for rel in self.doc.part.rels.values() 
                               if "image" in rel.target_ref.lower()])
        issues['tables'] = len(self.doc.tables)
        issues['headers'] = len(self.doc.sections[0].header.paragraphs) if self.doc.sections else 0
        issues['footers'] = len(self.doc.sections[0].footer.paragraphs) if self.doc.sections else 0
        
        # Extract text from main document paragraphs
        for para in self.doc.paragraphs:
            text = para.text.strip()
            if text:  # Only add non-empty paragraphs
                paragraphs.append(text)
        
        return paragraphs, issues


def process_document_set(
    doc_paths: Dict[str, Path],
    output_base: Path,
    discarded_base: Path,
    folder_name: str,
    doc_name: str
) -> None:
    """
    Process a set of 5 parallel documents (one per language).
    
    Args:
        doc_paths: Dictionary mapping language codes to document paths
        output_base: Base path for parallel output
        discarded_base: Base path for discarded content
        folder_name: Name of the source folder
        doc_name: Name of the document (without language suffix)
    """
    print("\n" + "="*60)
    print("Processing document set")
    print("="*60)
    
    # Extract paragraphs from all language versions
    all_paragraphs = {}
    all_issues = {}
    
    for lang_code, doc_path in sorted(doc_paths.items()):
        print(f"\nProcessing {lang_code}: {doc_path.name}")
        extractor = DocxParagraphExtractor(doc_path)
        paragraphs, issues = extractor.extract_paragraphs()
        all_paragraphs[lang_code] = paragraphs
        all_issues[lang_code] = issues
        
        print(f"  - Extracted {len(paragraphs)} paragraph(s)")
        
        # Report issues
        issue_messages = []
        if issues['images'] > 0:
            issue_messages.append(f"Found {issues['images']} image(s) - images ignored")
        if issues['tables'] > 0:
            issue_messages.append(f"Found {issues['tables']} table(s) - tables ignored")
        if issues['headers'] > 0:
            issue_messages.append(f"Found {issues['headers']} header(s) - headers ignored")
        if issues['footers'] > 0:
            issue_messages.append(f"Found {issues['footers']} footer(s) - footers ignored")
        
        if issue_messages:
            print("  - Issues found:")
            for msg in issue_messages:
                print(f"    * {msg}")
    
    # Compare paragraph counts
    print("\n" + "="*60)
    print("Initial paragraph count comparison")
    print("="*60)
    print()
    
    counts = {lang: len(paras) for lang, paras in all_paragraphs.items()}
    min_count = min(counts.values())
    max_count = max(counts.values())
    
    for lang_code in sorted(all_paragraphs.keys()):
        count = counts[lang_code]
        if count == max_count and count != min_count:
            symbol = "✓"
        elif count == min_count and count != max_count:
            symbol = "⚠"
        else:
            symbol = "•"
        print(f"{symbol} {lang_code}: {count} paragraph(s)")
    
    print()
    
    if min_count != max_count:
        print("⚠ Paragraph counts differ across languages!")
        print(f"  Range: {min_count} to {max_count} paragraph(s)")
        print(f"  Difference: {max_count - min_count} paragraph(s)")
        print()
        print(f"  → Keeping only the first {min_count} paragraph(s) from each version")
        print(f"  → Discarding {max_count - min_count} extra paragraph(s) from longer versions")
        
        # Save discarded paragraphs
        discarded_folder = discarded_base / doc_name
        for lang_code, paragraphs in all_paragraphs.items():
            if len(paragraphs) > min_count:
                discarded_folder.mkdir(parents=True, exist_ok=True)
                discarded_path = discarded_folder / f"{lang_code}.txt"
                discarded_content = '\n\n'.join(paragraphs[min_count:])
                discarded_path.write_text(discarded_content, encoding='utf-8')
                print(f"    • {lang_code}: discarded {len(paragraphs) - min_count} paragraph(s) → {discarded_path}")
    else:
        print(f"✓ All languages have the same number of paragraphs: {min_count}")
    
    # Trim all to minimum count
    for lang_code in all_paragraphs:
        all_paragraphs[lang_code] = all_paragraphs[lang_code][:min_count]
    
    # Save parallel paragraphs
    print("\n" + "="*60)
    print("Saving parallel paragraphs")
    print("="*60)
    print()
    
    output_folder = output_base / folder_name / doc_name
    output_folder.mkdir(parents=True, exist_ok=True)
    
    for lang_code, paragraphs in sorted(all_paragraphs.items()):
        output_path = output_folder / f"{lang_code}.txt"
        # Join paragraphs with double newline
        content = '\n\n'.join(paragraphs)
        output_path.write_text(content, encoding='utf-8')
        print(f"Saved {lang_code}: {output_path} ({len(paragraphs)} paragraph(s))")
    
    print()
    print(f"✓ All output files now contain {min_count} parallel paragraph(s)")


def discover_document_sets(base_path: Path) -> Dict[str, Dict[str, Dict[str, Path]]]:
    """
    Discover all complete document sets in the folder structure.
    
    Returns:
        Dictionary: {folder_name: {doc_name: {lang_code: path}}}
    """
    folders = {
        'Presidente': base_path / 'Presidente',
        'Consejo de Ministros': base_path / 'Consejo de Ministros',
        'Gobierno XV': base_path / 'Gobierno XV',
        'Actualidad': base_path / 'Actualidad'
    }
    
    language_patterns = {
        'castellano': r'[_-]castellano\.docx$',
        'ca_ES': r'[_-](ca-ES|ca_ES)\.docx$',
        'vl_ES': r'[_-](ca-ES-valencia|va_ES|vl_ES)\.docx$',
        'gl_ES': r'[_-](gl-ES|gl_ES|ga_ES)\.docx$',
        'eu_ES': r'[_-](eu-ES|eu_ES)\.docx$'
    }
    
    all_sets = {}
    
    for folder_name, folder_path in folders.items():
        if not folder_path.exists():
            continue
        
        # Group files by base name
        file_groups = {}
        
        for docx_file in folder_path.glob('*.docx'):
            if docx_file.name.startswith('~'):  # Skip temp files
                continue
            
            # Determine language and base name
            lang_found = None
            base_name = None
            
            for lang_code, pattern in language_patterns.items():
                if re.search(pattern, docx_file.name):
                    lang_found = lang_code
                    base_name = re.sub(pattern, '', docx_file.name)
                    break
            
            if lang_found and base_name:
                if base_name not in file_groups:
                    file_groups[base_name] = {}
                file_groups[base_name][lang_found] = docx_file
        
        # Keep only complete sets (all 5 languages)
        complete_sets = {}
        for base_name, lang_files in file_groups.items():
            if len(lang_files) == 5:
                complete_sets[base_name] = lang_files
        
        if complete_sets:
            all_sets[folder_name] = complete_sets
    
    return all_sets


def main():
    """Main execution function."""
    base_path = Path(__file__).parent / 'original_texts'
    output_base = Path(__file__).parent / 'extracted_paragraphs'
    discarded_base = Path(__file__).parent / 'extracted_paragraphs' / 'discarded_texts'
    
    # Discover all document sets
    all_sets = discover_document_sets(base_path)
    
    # Count total documents
    total_docs = sum(len(docs) for docs in all_sets.values())
    
    print(f"\nFound {total_docs} complete document sets across {len(all_sets)} folders\n")
    
    for folder_name, doc_sets in all_sets.items():
        print(f"{folder_name}: {len(doc_sets)} document(s)")
        for doc_name in sorted(doc_sets.keys()):
            print(f"  - {doc_name}")
    
    print("\n" + "="*80)
    print("PROCESSING ALL DOCUMENTS")
    print("="*80)
    
    # Process all documents
    doc_counter = 0
    successful = 0
    failed = 0
    
    for folder_name in sorted(all_sets.keys()):
        doc_sets = all_sets[folder_name]
        
        for doc_name in sorted(doc_sets.keys()):
            doc_counter += 1
            print(f"\n\n{'#'*80}")
            print(f"[{doc_counter}/{total_docs}] {folder_name} > {doc_name}")
            print('#'*80)
            
            try:
                process_document_set(
                    doc_paths=doc_sets[doc_name],
                    output_base=output_base,
                    discarded_base=discarded_base,
                    folder_name=folder_name,
                    doc_name=doc_name
                )
                successful += 1
            except Exception as e:
                print(f"\n❌ Error processing {doc_name}: {e}")
                failed += 1
    
    print("\n\n" + "="*80)
    print("PROCESSING COMPLETE")
    print("="*80)
    print()
    print(f"Total documents: {total_docs}")
    print(f"Successfully processed: {successful}")
    print(f"Failed: {failed}")
    print()
    
    if failed == 0:
        print("✅ All documents processed successfully!")
    else:
        print(f"⚠ {failed} document(s) failed to process")
    
    print()
    print("Output directories:")
    print(f"  - Parallel texts: {output_base}")
    print(f"  - Discarded texts: {discarded_base}")


if __name__ == '__main__':
    main()
